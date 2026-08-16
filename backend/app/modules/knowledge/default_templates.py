"""内置通用默认 docx 模板（通用版式）"""
from __future__ import annotations

import io
import os
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.platform.config import KNOWLEDGE_GENERIC_TEMPLATE_DIR

# kind → 中文文件名与说明
BUILTIN_TEMPLATE_CATALOG: dict[str, dict[str, str]] = {
    "iteration_report": {
        "filename": "通用迭代自动化测试报告模板.docx",
        "label": "迭代自动化测试报告",
        "description": "汇总多份自动化执行记录、缺陷统计与 AI 结论，适合迭代自动化收尾",
    },
    "functional_report": {
        "filename": "通用功能测试报告模板.docx",
        "label": "功能测试报告",
        "description": "侧重功能范围、缺陷与测试结论",
    },
    "performance_report": {
        "filename": "通用性能测试报告模板.docx",
        "label": "性能测试报告",
        "description": "侧重性能指标、瓶颈与结论",
    },
    "test_plan": {
        "filename": "通用测试计划模板.docx",
        "label": "测试计划",
        "description": "测试目标、范围、策略与环境",
    },
    "test_scheme": {
        "filename": "通用测试方案模板.docx",
        "label": "测试方案",
        "description": "方案级说明，含目标、策略与风险",
    },
}

# pptx 内置模板（通用质量回顾；程序按版式填充，非 docxtpl）
PPTX_BUILTIN_TEMPLATE_CATALOG: dict[str, dict[str, str]] = {
    "quality_review": {
        "filename": "通用质量回顾模板.pptx",
        "label": "质量回顾",
        "description": "通用版质量回顾 pptx（程序按版式填充）",
    },
}

DEFAULT_TEMPLATE_KIND = "iteration_report"
DEFAULT_TEMPLATE_FILENAME = BUILTIN_TEMPLATE_CATALOG[DEFAULT_TEMPLATE_KIND]["filename"]

_FONT = "微软雅黑"
_COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
_COLOR_SECONDARY = RGBColor(0x5B, 0x7C, 0x99)
_COLOR_MUTED = RGBColor(0x70, 0x70, 0x70)
_FILL_HEADER = "D6E4F0"
_FILL_ACCENT = "EEF4FA"


def _set_run_font(
    run,
    size_pt: float = 11,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    italic: bool = False,
) -> None:
    run.font.name = _FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size_pt: float = 10.5,
    align_center: bool = False,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if align_center else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    _set_run_font(run, size_pt, bold=bold, color=color)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)
    for level, size in ((1, 16), (2, 13)):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = _FONT
            h._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = _COLOR_PRIMARY
            h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
            h.paragraph_format.space_after = Pt(8)


def _new_document() -> Document:
    doc = Document()
    _configure_document(doc)
    return doc


def _cover_block(doc: Document, title: str, *, subtitle: str = "", badge: str = "") -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    _set_run_font(run, 22, bold=True, color=_COLOR_PRIMARY)

    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub.paragraph_format.space_after = Pt(4)
        run_sub = sub.add_run(subtitle)
        _set_run_font(run_sub, 15, color=_COLOR_SECONDARY)

    if badge:
        badge_p = doc.add_paragraph()
        badge_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        badge_p.paragraph_format.space_after = Pt(12)
        run_badge = badge_p.add_run(badge)
        _set_run_font(run_badge, 10, color=_COLOR_MUTED)

    line = doc.add_paragraph()
    line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line.paragraph_format.space_after = Pt(14)
    run_line = line.add_run("—" * 18)
    _set_run_font(run_line, 10, color=_COLOR_SECONDARY)


def _meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for idx, (label, value) in enumerate(rows):
        row = table.rows[idx]
        _shade_cell(row.cells[0], _FILL_HEADER)
        _set_cell_text(row.cells[0], label, bold=True, size_pt=10, color=_COLOR_PRIMARY)
        _set_cell_text(row.cells[1], value, size_pt=10.5)
    doc.add_paragraph()


def _meta_line(doc: Document) -> None:
    _meta_table(
        doc,
        [
            ("报告日期", "{{ report_date }}"),
            ("编写人", "{{ author }}"),
            ("测试人员", "{{ tester_name }}"),
        ],
    )


def _add_hint(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _set_run_font(run, 9.5, italic=True, color=_COLOR_MUTED)


def _add_placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, 11)


def _add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_multiline_block(doc: Document, loop_var: str) -> None:
    doc.add_paragraph(f"{{%p for line in {loop_var} %}}")
    _add_placeholder(doc, "{{ line }}")
    doc.add_paragraph("{%p endfor %}")


def _add_bullet_loop(doc: Document, loop_expr: str, item_template: str, *, title: str = "") -> None:
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        _set_run_font(run, 11, bold=True, color=_COLOR_PRIMARY)
    doc.add_paragraph(f"{{%p for {loop_expr} %}}")
    _add_placeholder(doc, item_template)
    doc.add_paragraph("{%p endfor %}")


def _add_stat_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=2, cols=len(items))
    table.style = "Table Grid"
    for idx, (label, value) in enumerate(items):
        _shade_cell(table.rows[0].cells[idx], _FILL_ACCENT)
        _set_cell_text(table.rows[0].cells[idx], label, bold=True, size_pt=10, align_center=True, color=_COLOR_PRIMARY)
        _set_cell_text(table.rows[1].cells[idx], value, size_pt=11, align_center=True)
    doc.add_paragraph()


def _add_table_loop(
    doc: Document,
    *,
    headers: list[str],
    loop_var: str,
    fields: list[str],
    loop_item: str = "row",
    title: str = "",
) -> None:
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        _set_run_font(run, 11, bold=True, color=_COLOR_PRIMARY)
    table = doc.add_table(rows=3, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        _shade_cell(table.rows[0].cells[idx], _FILL_HEADER)
        _set_cell_text(table.rows[0].cells[idx], header, bold=True, size_pt=10, align_center=True, color=_COLOR_PRIMARY)
    table.rows[1].cells[0].text = "{%tr for " + loop_item + " in " + loop_var + " %}"
    data = table.rows[2].cells
    for idx, field in enumerate(fields):
        data[idx].text = "{{ " + loop_item + "." + field + " }}"
    end_row = table.add_row()
    end_row.cells[0].text = "{%tr endfor %}"
    doc.add_paragraph()


def _bug_section(doc: Document) -> None:
    _add_stat_table(
        doc,
        [
            ("Bug 总数", "{{ bug_total }}"),
            ("未关闭", "{{ bug_open_count }}"),
            ("已关闭", "{{ bug_closed_count }}"),
        ],
    )
    _add_placeholder(doc, "{{ bug_export_summary }}")
    _add_table_loop(
        doc,
        title="缺陷分布",
        headers=["统计维度", "项", "数量"],
        loop_var="bug_stats_rows",
        fields=["dimension", "label", "count"],
    )
    _add_multiline_block(doc, "bug_analysis_lines")


def _appendix_summary_block(doc: Document) -> None:
    _add_section_heading(doc, "附录：资料摘要", level=1)
    _add_hint(doc, "以下为重点摘要，供评审追溯；完整文档见资料库文件夹。")
    _add_placeholder(doc, "{{ knowledge_summary }}")


def _source_refs_section(doc: Document) -> None:
    _add_section_heading(doc, "附录：报告数据来源", level=2)
    _add_hint(doc, "下列为本次报告生成时纳入的资料与模板引用。")
    _add_bullet_loop(
        doc,
        "ref in source_refs",
        "- {{ ref.kind_label }}：{{ ref.title }}（{{ ref.detail }}）",
    )


def build_iteration_report_template() -> bytes:
    doc = _new_document()
    _cover_block(
        doc,
        "{{ project_name }} · {{ iteration_name }}",
        subtitle="迭代自动化测试报告",
        badge="BrickCore · 内置模板",
    )
    _meta_line(doc)

    _add_section_heading(doc, "一、测试概述")
    _add_multiline_block(doc, "overview_lines")

    _add_section_heading(doc, "二、测试范围")
    _add_multiline_block(doc, "test_scope_lines")

    _add_section_heading(doc, "三、自动化执行情况")
    _add_stat_table(
        doc,
        [
            ("通过率", "{{ pass_rate_display }}"),
            ("用例总数", "{{ total_cases }}"),
            ("失败数", "{{ failed_case_count }}"),
            ("耗时", "{{ exec_duration }}"),
        ],
    )
    _add_bullet_loop(
        doc,
        "rec in exec_records",
        "- {{ rec.name }}（{{ rec.report_type }}）通过率 {{ rec.pass_rate }}%，共 {{ rec.total_cases }} 条",
        title="执行记录摘要",
    )
    _add_bullet_loop(
        doc,
        "item in failed_cases",
        "- {{ item.case_name }} | {{ item.status }} | {{ item.error_msg }}",
        title="失败用例",
    )

    _add_section_heading(doc, "四、缺陷情况")
    _bug_section(doc)

    _add_section_heading(doc, "五、测试结论与风险")
    _add_multiline_block(doc, "ai_conclusion_lines")
    _add_multiline_block(doc, "ai_risks_lines")

    _add_section_heading(doc, "六、改进建议")
    _add_multiline_block(doc, "ai_recommendations_lines")

    doc.add_page_break()
    _appendix_summary_block(doc)
    _source_refs_section(doc)

    return _doc_to_bytes(doc)


def build_functional_report_template() -> bytes:
    doc = _new_document()
    _cover_block(doc, "{{ project_name }} · {{ iteration_name }}", subtitle="功能测试报告", badge="BrickCore · 内置模板")
    _meta_line(doc)

    _add_section_heading(doc, "一、测试概述")
    _add_multiline_block(doc, "overview_lines")

    _add_section_heading(doc, "二、功能测试范围")
    _add_multiline_block(doc, "functional_scope_lines")
    _add_multiline_block(doc, "test_scope_lines")

    _add_section_heading(doc, "三、测试环境与人员")
    _meta_table(
        doc,
        [
            ("测试环境", "{{ test_environment }}"),
            ("审批人", "{{ approver }}"),
        ],
    )

    _add_section_heading(doc, "四、功能测试执行情况")
    _add_stat_table(
        doc,
        [
            ("用例总数", "{{ total_cases }}"),
            ("失败数", "{{ failed_case_count }}"),
            ("通过率", "{{ pass_rate_display }}"),
        ],
    )
    _add_bullet_loop(
        doc,
        "item in failed_cases",
        "- {{ item.case_name }} | {{ item.error_msg }}",
        title="失败用例",
    )

    _add_section_heading(doc, "五、缺陷分析")
    _bug_section(doc)

    _add_section_heading(doc, "六、功能测试总结")
    _add_multiline_block(doc, "functional_summary_lines")
    _add_multiline_block(doc, "ai_conclusion_lines")

    _add_section_heading(doc, "七、遗留风险与建议")
    _add_multiline_block(doc, "ai_risks_lines")
    _add_multiline_block(doc, "ai_recommendations_lines")

    doc.add_page_break()
    _source_refs_section(doc)

    return _doc_to_bytes(doc)


def build_performance_report_template() -> bytes:
    doc = _new_document()
    _cover_block(doc, "{{ project_name }} · {{ iteration_name }}", subtitle="性能测试报告", badge="BrickCore · 内置模板")
    _meta_line(doc)

    _add_section_heading(doc, "一、测试概述")
    _add_multiline_block(doc, "overview_lines")

    _add_section_heading(doc, "二、测试范围与目标")
    _add_multiline_block(doc, "test_scope_lines")

    _add_section_heading(doc, "三、测试环境")
    _add_placeholder(doc, "{{ test_environment }}")

    _add_section_heading(doc, "四、性能指标与结果")
    _add_placeholder(doc, "{{ performance_metrics }}")
    _add_stat_table(doc, [("执行耗时", "{{ exec_duration }}"), ("Bug 总数", "{{ bug_total }}")])

    _add_section_heading(doc, "五、性能测试总结")
    _add_placeholder(doc, "{{ performance_summary }}")
    _add_placeholder(doc, "{{ ai_conclusion }}")

    _add_section_heading(doc, "六、缺陷与风险")
    _add_placeholder(doc, "{{ bug_export_summary }}")
    _add_placeholder(doc, "{{ ai_risks }}")

    _add_section_heading(doc, "七、优化建议")
    _add_placeholder(doc, "{{ ai_recommendations }}")

    doc.add_page_break()
    _source_refs_section(doc)

    return _doc_to_bytes(doc)


def build_test_plan_template() -> bytes:
    doc = _new_document()
    _cover_block(doc, "{{ project_name }} · {{ iteration_name }}", subtitle="测试计划", badge="BrickCore · 内置模板")
    _meta_line(doc)

    _add_section_heading(doc, "一、项目背景")
    _add_placeholder(doc, "{{ overview }}")

    _add_section_heading(doc, "二、测试目标")
    _add_placeholder(doc, "{{ test_objectives }}")

    _add_section_heading(doc, "三、测试范围")
    _add_placeholder(doc, "{{ test_scope }}")

    _add_section_heading(doc, "四、测试策略")
    _add_placeholder(doc, "{{ test_strategy }}")

    _add_section_heading(doc, "五、测试环境与资源")
    _meta_table(
        doc,
        [
            ("测试环境", "{{ test_environment }}"),
            ("测试人员", "{{ tester_name }}"),
            ("审批人", "{{ approver }}"),
        ],
    )

    _add_section_heading(doc, "六、进度与里程碑")
    _meta_table(
        doc,
        [
            ("计划开始", "{{ plan_start_date }}"),
            ("计划结束", "{{ plan_end_date }}"),
            ("总负责人", "{{ plan_owner }}"),
        ],
    )
    _add_table_loop(
        doc,
        title="里程碑安排",
        headers=["里程碑", "开始日期", "结束日期", "负责人", "说明"],
        loop_var="milestone_rows",
        fields=["name", "start_date", "end_date", "owner", "remark"],
    )

    _add_section_heading(doc, "七、风险与应对")
    _add_placeholder(doc, "{{ ai_risks }}")

    doc.add_page_break()
    _appendix_summary_block(doc)
    _source_refs_section(doc)

    return _doc_to_bytes(doc)


def build_test_scheme_template() -> bytes:
    doc = _new_document()
    _cover_block(doc, "{{ project_name }} · {{ iteration_name }}", subtitle="测试方案", badge="BrickCore · 内置模板")
    _meta_line(doc)

    _add_section_heading(doc, "一、引言")
    _add_placeholder(doc, "{{ overview }}")

    _add_section_heading(doc, "二、测试目标")
    _add_placeholder(doc, "{{ test_objectives }}")

    _add_section_heading(doc, "三、测试范围")
    _add_placeholder(doc, "{{ test_scope }}")

    _add_section_heading(doc, "四、测试类型与策略")
    _meta_table(
        doc,
        [
            ("功能测试", "{{ functional_scope }}"),
            ("性能测试", "{{ performance_metrics }}"),
            ("总体策略", "{{ test_strategy }}"),
        ],
    )

    _add_section_heading(doc, "五、测试组织与职责")
    _meta_table(
        doc,
        [
            ("测试负责人", "{{ tester_name }}"),
            ("审批人", "{{ approver }}"),
        ],
    )

    _add_section_heading(doc, "六、测试环境")
    _add_placeholder(doc, "{{ test_environment }}")

    _add_section_heading(doc, "七、准入准出标准")
    _add_hint(doc, "可结合通过率、缺陷数等指标定义准入准出；下列为当前资料库参考值。")
    _add_stat_table(
        doc,
        [
            ("通过率(%)", "{{ pass_rate }}"),
            ("Bug 总数", "{{ bug_total }}"),
        ],
    )

    _add_section_heading(doc, "八、风险分析")
    _add_placeholder(doc, "{{ ai_risks }}")

    doc.add_page_break()
    _appendix_summary_block(doc)
    _source_refs_section(doc)

    return _doc_to_bytes(doc)


_BUILDERS: dict[str, Callable[[], bytes]] = {
    "iteration_report": build_iteration_report_template,
    "functional_report": build_functional_report_template,
    "performance_report": build_performance_report_template,
    "test_plan": build_test_plan_template,
    "test_scheme": build_test_scheme_template,
}


def build_generic_quality_review_pptx_template() -> bytes:
    """通用质量回顾内置模板（程序化生成）。"""
    try:
        from app.modules.knowledge.packs.digitech.quality_pptx_builder import build_quality_review_pptx
    except ImportError as ex:
        raise RuntimeError("质量回顾 pptx 生成依赖行业扩展包，当前环境未安装") from ex

    return build_quality_review_pptx(
        {
            "scheme_name": "示例迭代",
            "iteration_no": "2026S1",
            "author": "测试负责人",
            "report_date": "2026-01-01",
            "quality_summary": "（通用版：上传 Bug 导出与任务/工时单后自动生成）",
            "iteration_summary": "本迭代开发耗时 — 个工时，共发现 bug 数 — 个…",
            "d1_score": 0,
            "d1_grade": "-",
            "reporter_rows": [],
            "developer_rows": [],
            "release_issue_rows": [],
        }
    )


_PPTX_BUILDERS: dict[str, Callable[[], bytes]] = {
    "quality_review": build_generic_quality_review_pptx_template,
}


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def is_pptx_builtin_kind(kind: str) -> bool:
    return kind in PPTX_BUILTIN_TEMPLATE_CATALOG


def _pptx_disk_candidates(kind: str) -> list[str]:
    meta = PPTX_BUILTIN_TEMPLATE_CATALOG[kind]
    names: list[str] = []
    if meta.get("filename"):
        names.append(meta["filename"])
    return names


def list_builtin_templates() -> list[dict[str, str]]:
    items = [{"kind": kind, "format": "docx", **meta} for kind, meta in BUILTIN_TEMPLATE_CATALOG.items()]
    # 质量回顾 pptx 依赖行业扩展包；未开通时不出现在列表，避免下载 500
    try:
        from app.core.platform.edition import knowledge_pack_addon_enabled
    except Exception:  # pragma: no cover
        knowledge_pack_addon_enabled = lambda: True  # type: ignore[assignment]
    if knowledge_pack_addon_enabled():
        items.extend(
            {"kind": kind, "format": "pptx", **meta}
            for kind, meta in PPTX_BUILTIN_TEMPLATE_CATALOG.items()
        )
    return items


def build_builtin_template(kind: str) -> tuple[bytes, str]:
    if kind in _BUILDERS:
        meta = BUILTIN_TEMPLATE_CATALOG[kind]
        return _BUILDERS[kind](), meta["filename"]
    if is_pptx_builtin_kind(kind):
        return load_builtin_template_bytes(kind)
    raise ValueError(f"未知内置模板类型: {kind}")


# 兼容旧测试/调用
def build_generic_iteration_report_template() -> bytes:
    return build_iteration_report_template()


def _read_disk_template(kind: str) -> tuple[bytes, str] | None:
    if is_pptx_builtin_kind(kind):
        meta = PPTX_BUILTIN_TEMPLATE_CATALOG[kind]
        for name in _pptx_disk_candidates(kind):
            path = os.path.join(KNOWLEDGE_GENERIC_TEMPLATE_DIR, name)
            if os.path.isfile(path):
                with open(path, "rb") as f:
                    return f.read(), meta["filename"]
        return None

    meta = BUILTIN_TEMPLATE_CATALOG.get(kind)
    if not meta:
        return None
    path = os.path.join(KNOWLEDGE_GENERIC_TEMPLATE_DIR, meta["filename"])
    if os.path.isfile(path):
        with open(path, "rb") as f:
            return f.read(), meta["filename"]
    return None


def ensure_builtin_templates_on_disk() -> None:
    """仅当磁盘无对应文件时生成 docx 默认模板；已有 Git 跟踪文件不覆盖。"""
    os.makedirs(KNOWLEDGE_GENERIC_TEMPLATE_DIR, exist_ok=True)
    for kind in _BUILDERS:
        meta = BUILTIN_TEMPLATE_CATALOG[kind]
        path = os.path.join(KNOWLEDGE_GENERIC_TEMPLATE_DIR, meta["filename"])
        if os.path.isfile(path):
            continue
        content, _ = _BUILDERS[kind](), meta["filename"]
        with open(path, "wb") as f:
            f.write(content)


def ensure_generic_template_on_disk() -> str:
    ensure_builtin_templates_on_disk()
    return os.path.join(KNOWLEDGE_GENERIC_TEMPLATE_DIR, DEFAULT_TEMPLATE_FILENAME)


def load_builtin_template_bytes(kind: str = DEFAULT_TEMPLATE_KIND) -> tuple[bytes, str]:
    ensure_builtin_templates_on_disk()
    on_disk = _read_disk_template(kind)
    if on_disk:
        return on_disk
    if kind in _PPTX_BUILDERS:
        meta = PPTX_BUILTIN_TEMPLATE_CATALOG[kind]
        return _PPTX_BUILDERS[kind](), meta["filename"]
    if kind in _BUILDERS:
        meta = BUILTIN_TEMPLATE_CATALOG[kind]
        return _BUILDERS[kind](), meta["filename"]
    if is_pptx_builtin_kind(kind):
        raise FileNotFoundError(
            f"未找到质量回顾内置模板，请将 pptx 置于 {KNOWLEDGE_GENERIC_TEMPLATE_DIR}"
        )
    raise ValueError(f"未知内置模板类型: {kind}")


def load_generic_template_bytes() -> bytes:
    content, _ = load_builtin_template_bytes(DEFAULT_TEMPLATE_KIND)
    return content
