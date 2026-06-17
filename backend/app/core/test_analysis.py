"""
AI 测试分析：测试点归一化、思维导图树、XMind 导出、测试方案渲染
"""
from __future__ import annotations

import io
import json
import re
import uuid
import zipfile
from collections import defaultdict
from typing import Any, Optional

VALID_TEST_TYPES = {"正向", "异常", "边界", "权限", "状态", "兼容", "性能", "API", "其他"}
PRIORITY_MAP = {"P0": "P0", "P1": "P1", "P2": "P2", "P3": "P3", "1": "P0", "2": "P1", "3": "P2", "4": "P3"}


def _as_text(val) -> str:
    if val is None:
        return ""
    if isinstance(val, list):
        return "\n".join(str(v) for v in val if v)
    return str(val).strip()


def _new_id() -> str:
    return uuid.uuid4().hex[:16]


def normalize_test_point(item: dict, *, section_ids: Optional[list] = None) -> Optional[dict]:
    if not isinstance(item, dict):
        return None
    title = (item.get("title") or item.get("name") or "").strip()
    if not title:
        return None
    test_type = (item.get("test_type") or item.get("type") or "正向").strip()
    if test_type not in VALID_TEST_TYPES:
        test_type = "其他" if test_type else "正向"
    priority = str(item.get("priority") or "P2").strip().upper()
    if priority.startswith("P"):
        priority = PRIORITY_MAP.get(priority, "P2")
    else:
        priority = PRIORITY_MAP.get(priority, "P2")
    main_module = (item.get("main_module") or "").strip()
    sub_module = (item.get("sub_module") or "").strip()
    module_path = (item.get("module_path") or item.get("module") or "").strip()
    if not module_path and (main_module or sub_module):
        module_path = "/".join(p for p in [main_module, sub_module] if p)
    if not main_module and module_path:
        parts = [p.strip() for p in module_path.split("/") if p.strip()]
        main_module = parts[0] if parts else "未分类"
        sub_module = parts[1] if len(parts) > 1 else (sub_module or "通用")
    if not main_module:
        main_module = "未分类"
    if not sub_module:
        sub_module = "通用"
    return {
        "title": title[:300],
        "description": (item.get("description") or item.get("desc") or "").strip()[:5000],
        "test_type": test_type,
        "priority": priority,
        "module_path": module_path[:200],
        "main_module": main_module[:100],
        "sub_module": sub_module[:100],
        "acceptance_ref": _as_text(item.get("acceptance_ref") or item.get("acceptance_criteria"))[:2000],
        "section_ids": section_ids or item.get("section_ids") or [],
        "extra": {
            "tags": item.get("tags") or [],
            "feature_point": (item.get("feature_point") or "").strip(),
        },
    }


def normalize_test_points(raw_items: list, *, section_ids: Optional[list] = None) -> list[dict]:
    out: list[dict] = []
    seen: set[str] = set()
    for item in raw_items or []:
        norm = normalize_test_point(item, section_ids=section_ids)
        if not norm:
            continue
        key = f"{norm['main_module']}|{norm['sub_module']}|{norm['title']}"
        if key in seen:
            continue
        seen.add(key)
        out.append(norm)
    return out


def _point_label(point: dict) -> str:
    parts = []
    if point.get("priority"):
        parts.append(f"({point['priority']})")
    if point.get("test_type"):
        parts.append(f"[{point['test_type']}]")
    parts.append(point.get("title") or "")
    return " ".join(parts)


def filter_test_points(
    points: list[dict],
    *,
    section_ids: Optional[list[str]] = None,
    batch_name: Optional[str] = None,
    status: Optional[str] = None,
    test_type: Optional[str] = None,
    priority: Optional[str] = None,
    main_module: Optional[str] = None,
    sub_module: Optional[str] = None,
    keyword: Optional[str] = None,
) -> list[dict]:
    out = list(points or [])
    if batch_name and batch_name.strip():
        ref = f"batch:{batch_name.strip()}"
        out = [p for p in out if (p.get("source_ref") or "") == ref]
    if section_ids:
        id_set = set(section_ids)
        out = [
            p for p in out
            if id_set.intersection(set(p.get("section_ids") or []))
        ]
    if status:
        out = [p for p in out if (p.get("status") or "") == status]
    if test_type:
        out = [p for p in out if (p.get("test_type") or "") == test_type]
    if priority:
        out = [p for p in out if (p.get("priority") or "") == priority]
    if main_module and main_module.strip():
        kw = main_module.strip()
        out = [p for p in out if (p.get("main_module") or "").strip() == kw]
    if sub_module and sub_module.strip():
        kw = sub_module.strip()
        out = [p for p in out if (p.get("sub_module") or "").strip() == kw]
    if keyword and keyword.strip():
        kw = keyword.strip().lower()
        out = [
            p for p in out
            if kw in (p.get("title") or "").lower()
            or kw in (p.get("description") or "").lower()
            or kw in (p.get("sub_module") or "").lower()
            or kw in (p.get("main_module") or "").lower()
        ]
    return out


def sort_test_points_list(points: list[dict]) -> list[dict]:
    """按主模块、子模块、sort_order、ID 排序，便于列表与导图展示。"""
    return sorted(
        points or [],
        key=lambda p: (
            (p.get("main_module") or "").strip(),
            (p.get("sub_module") or "").strip(),
            p.get("sort_order") or 0,
            p.get("id") or 0,
        ),
    )


def _section_depth(section_id: str, section_by_id: dict[str, dict]) -> int:
    depth = 0
    cur = section_by_id.get(section_id)
    visited: set[str] = set()
    while cur and cur.get("id") not in visited:
        visited.add(cur["id"])
        depth += 1
        pid = cur.get("parent_id")
        cur = section_by_id.get(pid) if pid else None
    return depth


def map_point_section_titles(
    item: dict,
    all_sections: list[dict],
) -> list[str]:
    """将 LLM 返回的 section_title / 子模块名映射为 section_ids（取最具体章节）。"""
    sections = [s for s in all_sections if s.get("id")]
    section_by_id = {s["id"]: s for s in sections}
    title_to_id = {
        (s.get("title") or "").strip(): s["id"]
        for s in sections
        if (s.get("title") or "").strip()
    }
    if not title_to_id:
        return item.get("section_ids") or []

    ref = (item.get("section_title") or item.get("section_ref") or "").strip()
    if ref and ref in title_to_id:
        return [title_to_id[ref]]

    sub = (item.get("sub_module") or "").strip()
    main = (item.get("main_module") or "").strip()
    matched: list[str] = []
    for title, sid in title_to_id.items():
        if sub and (sub in title or title in sub):
            matched.append(sid)
        elif not sub and main and (main in title or title in main):
            matched.append(sid)
    if matched:
        return [max(matched, key=lambda sid: _section_depth(sid, section_by_id))]
    return item.get("section_ids") or []


def _pick_section_for_point(
    point: dict,
    section_by_id: dict[str, dict],
    sections: Optional[list[dict]] = None,
) -> Optional[str]:
    sub = (point.get("sub_module") or "").strip()
    main = (point.get("main_module") or "").strip()
    if sections and (sub or main):
        matched: list[str] = []
        for s in sections:
            title = (s.get("title") or "").strip()
            sid = s.get("id")
            if not sid or sid not in section_by_id:
                continue
            if sub and (sub in title or title in sub):
                matched.append(sid)
            elif not sub and main and (main in title or title in main):
                matched.append(sid)
        if matched:
            return max(matched, key=lambda sid: _section_depth(sid, section_by_id))

    ids = [i for i in (point.get("section_ids") or []) if i in section_by_id]
    if not ids:
        return None
    return max(ids, key=lambda sid: _section_depth(sid, section_by_id))


def build_mindmap_tree_by_sections(
    root_title: str,
    points: list[dict],
    sections: list[dict],
) -> dict:
    """按需求章节层级展示测试点（优先于 LLM 模块名）。"""
    section_by_id = {s["id"]: s for s in sections if s.get("id")}
    if not section_by_id:
        return build_mindmap_tree(root_title, points)

    section_points: dict[str, list[dict]] = defaultdict(list)
    orphans: list[dict] = []
    for idx, pt in enumerate(points):
        pt = {**pt, "sort_order": idx}
        sid = _pick_section_for_point(pt, section_by_id, sections)
        if sid:
            section_points[sid].append(pt)
        else:
            orphans.append(pt)

    used_ids: set[str] = set()

    def collect_ancestors(sec_id: str) -> None:
        cur = section_by_id.get(sec_id)
        while cur:
            used_ids.add(cur["id"])
            pid = cur.get("parent_id")
            cur = section_by_id.get(pid) if pid else None

    for sid in section_points:
        collect_ancestors(sid)

    def build_section_node(sec_id: str) -> dict:
        sec = section_by_id[sec_id]
        children: list[dict] = []
        child_secs = [
            s for s in sections
            if s.get("parent_id") == sec_id and s.get("id") in used_ids
        ]
        child_secs.sort(key=lambda s: (s.get("level") or 0, s.get("title") or ""))
        for cs in child_secs:
            children.append(build_section_node(cs["id"]))
        leaves = section_points.get(sec_id, [])
        leaves.sort(key=lambda x: (x.get("sort_order", 0), x.get("title") or ""))
        for pt in leaves:
            children.append({
                "name": _point_label(pt),
                "node_type": "test_point",
                "point_id": pt.get("id"),
                "title": pt.get("title"),
                "test_type": pt.get("test_type"),
                "priority": pt.get("priority"),
                "description": pt.get("description"),
                "status": pt.get("status"),
                "sort_order": pt.get("sort_order", 0),
            })
        return {
            "name": sec.get("title") or sec_id,
            "node_type": "section",
            "section_id": sec_id,
            "children": children,
        }

    root_children: list[dict] = []
    roots = [
        s for s in sections
        if s.get("id") in used_ids and (not s.get("parent_id") or s.get("parent_id") not in used_ids)
    ]
    if not roots:
        roots = [s for s in sections if s.get("id") in used_ids]
    roots.sort(key=lambda s: (s.get("level") or 0, s.get("title") or ""))
    for r in roots:
        root_children.append(build_section_node(r["id"]))

    if orphans:
        orphan_node = {
            "name": "未匹配章节",
            "node_type": "module",
            "children": [{
                "name": _point_label(pt),
                "node_type": "test_point",
                "point_id": pt.get("id"),
                "title": pt.get("title"),
                "test_type": pt.get("test_type"),
                "priority": pt.get("priority"),
                "description": pt.get("description"),
                "status": pt.get("status"),
                "sort_order": pt.get("sort_order", 0),
            } for pt in orphans],
        }
        root_children.append(orphan_node)

    return {
        "name": root_title or "测试点",
        "node_type": "root",
        "children": root_children,
    }


def resolve_mindmap_tree(
    root_title: str,
    points: list[dict],
    sections: Optional[list[dict]] = None,
    *,
    layout: str = "section",
) -> dict:
    if layout == "section" and sections:
        return build_mindmap_tree_by_sections(root_title, points, sections)
    return build_mindmap_tree(root_title, points)


def format_scope_section_titles(sections: list[dict]) -> str:
    if not sections:
        return "（未指定章节）"
    lines = []
    for s in sections:
        level = s.get("level") or 1
        indent = "  " * max(0, int(level) - 1)
        lines.append(f"{indent}- {s.get('title') or s.get('id')}")
    return "\n".join(lines)


def build_mindmap_tree(root_title: str, points: list[dict]) -> dict:
    """构建 ECharts tree 与通用嵌套树结构。"""
    root: dict[str, Any] = {
        "name": root_title or "测试点",
        "node_type": "root",
        "children": [],
    }
    main_map: dict[str, dict] = {}

    for idx, pt in enumerate(points):
        main = pt.get("main_module") or "未分类"
        sub = pt.get("sub_module") or "通用"
        if main not in main_map:
            main_map[main] = {"name": main, "node_type": "module", "children": {}}
        sub_map = main_map[main]["children"]
        if sub not in sub_map:
            sub_map[sub] = {"name": sub, "node_type": "sub_module", "children": []}
        leaf = {
            "name": _point_label(pt),
            "node_type": "test_point",
            "point_id": pt.get("id"),
            "title": pt.get("title"),
            "test_type": pt.get("test_type"),
            "priority": pt.get("priority"),
            "description": pt.get("description"),
            "status": pt.get("status"),
            "module_path": pt.get("module_path"),
            "sort_order": idx,
        }
        sub_map[sub]["children"].append(leaf)

    for main_name in sorted(main_map.keys()):
        main_node = main_map[main_name]
        main_children = []
        for sub_name in sorted(main_node["children"].keys()):
            sub_node = main_node["children"][sub_name]
            sub_node["children"].sort(key=lambda x: (x.get("sort_order", 0), x.get("title") or ""))
            main_children.append({
                "name": sub_node["name"],
                "node_type": "sub_module",
                "children": sub_node["children"],
            })
        root["children"].append({
            "name": main_name,
            "node_type": "module",
            "children": main_children,
        })
    return root


def _echarts_node(node: dict) -> dict:
    out: dict[str, Any] = {"name": node.get("name") or ""}
    meta = {
        "node_type": node.get("node_type"),
        "point_id": node.get("point_id"),
        "title": node.get("title"),
        "test_type": node.get("test_type"),
        "priority": node.get("priority"),
        "description": node.get("description"),
        "status": node.get("status"),
    }
    if any(v is not None for v in meta.values()):
        out["value"] = meta
    children = node.get("children") or []
    if children:
        out["children"] = [_echarts_node(c) for c in children]
    return out


def build_echarts_tree(
    root_title: str,
    points: list[dict],
    sections: Optional[list[dict]] = None,
    *,
    layout: str = "section",
) -> dict:
    tree = resolve_mindmap_tree(root_title, points, sections, layout=layout)
    return _echarts_node(tree)


def _topic_from_point(point: dict) -> dict:
    notes = point.get("description") or ""
    if point.get("acceptance_ref"):
        notes = f"{notes}\n\n验收依据：{point['acceptance_ref']}".strip()
    topic: dict[str, Any] = {
        "id": _new_id(),
        "class": "topic",
        "title": _point_label(point),
        "children": {"attached": []},
    }
    if notes:
        topic["notes"] = {"plain": {"content": notes[:8000]}}
    return topic


def _topic_from_node(node: dict) -> dict:
    children = node.get("children") or []
    attached = []
    for child in children:
        if child.get("node_type") == "test_point":
            attached.append(_topic_from_point({
                "title": child.get("title") or child.get("name"),
                "test_type": child.get("test_type"),
                "priority": child.get("priority"),
                "description": child.get("description"),
            }))
        else:
            attached.append(_topic_from_node(child))
    title = node.get("name") or ""
    topic: dict[str, Any] = {
        "id": _new_id(),
        "class": "topic",
        "title": title,
        "children": {"attached": attached},
    }
    if node.get("node_type") in (None, "root", "section", "module"):
        topic["structureClass"] = "org.xmind.ui.map.unbalanced"
    return topic


def build_xmind_bytes(root_title: str, points: list[dict], tree_root: Optional[dict] = None) -> bytes:
    """生成 XMind Zen 格式 (.xmind) ZIP 文件。"""
    tree = tree_root or build_mindmap_tree(root_title, points)
    root_topic = _topic_from_node(tree)
    content = [{
        "id": _new_id(),
        "revisionId": _new_id(),
        "class": "sheet",
        "title": "测试点",
        "rootTopic": root_topic,
        "topicPositioning": "fixed",
    }]
    metadata = {
        "creator": {"name": "BrickCore AI Test Analysis", "version": "1.0"},
        "activeSheetId": content[0]["id"],
    }
    manifest = {
        "file-entries": {
            "content.json": {},
            "metadata.json": {},
        }
    }
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("content.json", json.dumps(content, ensure_ascii=False, indent=2))
        zf.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False, indent=2))
        zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
    return buf.getvalue()


_POINT_TITLE_RE = re.compile(r"^\((P[0-3])\)\s*(?:\[([^\]]+)\]\s*)?(.*)$", re.I)


def _extract_xmind_topic_notes(topic: dict) -> str:
    notes = topic.get("notes") or {}
    if not isinstance(notes, dict):
        return _as_text(notes)
    plain = notes.get("plain")
    if isinstance(plain, dict):
        return (plain.get("content") or "").strip()
    if isinstance(plain, str):
        return plain.strip()
    real = notes.get("realHTML") or notes.get("html")
    if isinstance(real, dict):
        return _as_text(real.get("content"))
    return _as_text(real)


def _parse_xmind_leaf_title(raw: str) -> tuple[str, str, str]:
    text = (raw or "").strip()
    if not text:
        return "P2", "正向", ""
    match = _POINT_TITLE_RE.match(text)
    if match:
        priority = PRIORITY_MAP.get(match.group(1).upper(), "P2")
        test_type = (match.group(2) or "正向").strip()
        if test_type not in VALID_TEST_TYPES:
            test_type = "其他"
        title = (match.group(3) or text).strip()
        return priority, test_type, title
    return "P2", "正向", text


def _walk_xmind_topic(
    topic: dict,
    ancestors: list[str],
    out: list[dict],
    sort_order: int,
    *,
    is_root: bool = False,
) -> int:
    if not isinstance(topic, dict):
        return sort_order
    attached = (topic.get("children") or {}).get("attached") or []
    title = (topic.get("title") or "").strip()

    if not attached:
        if not title:
            return sort_order
        priority, test_type, clean_title = _parse_xmind_leaf_title(title)
        if not clean_title:
            return sort_order
        main = ancestors[0] if len(ancestors) >= 1 else "未分类"
        sub = ancestors[1] if len(ancestors) >= 2 else "通用"
        if len(ancestors) > 2:
            sub = "/".join(ancestors[1:])
        out.append({
            "title": clean_title,
            "description": _extract_xmind_topic_notes(topic)[:5000],
            "test_type": test_type,
            "priority": priority,
            "main_module": main[:100],
            "sub_module": sub[:100],
            "module_path": "/".join(p for p in [main, sub] if p)[:200],
            "sort_order": sort_order,
        })
        return sort_order + 1

    child_ancestors = list(ancestors)
    if title and not is_root:
        child_ancestors = ancestors + [title]
    idx = sort_order
    for child in attached:
        idx = _walk_xmind_topic(child, child_ancestors, out, idx, is_root=False)
    return idx


def _load_xmind_content_json(content: bytes) -> list[dict]:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            names = zf.namelist()
            target = "content.json" if "content.json" in names else None
            if not target:
                for name in names:
                    if name.endswith("content.json"):
                        target = name
                        break
            if not target:
                raise ValueError("未找到 content.json，请使用 XMind Zen 格式 (.xmind)")
            raw = zf.read(target).decode("utf-8")
    except zipfile.BadZipFile as e:
        raise ValueError("无效的 XMind 文件（需为 .xmind ZIP 格式）") from e
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError("XMind content.json 解析失败") from e
    if isinstance(data, list):
        return [s for s in data if isinstance(s, dict)]
    if isinstance(data, dict):
        return [data]
    raise ValueError("不支持的 XMind 内容结构")


def parse_xmind_to_test_points(content: bytes) -> list[dict]:
    """解析 XMind Zen (.xmind) 为测试点草稿列表。"""
    sheets = _load_xmind_content_json(content)
    raw_items: list[dict] = []
    sort_order = 0
    for sheet in sheets:
        root = sheet.get("rootTopic") or sheet.get("rootTopic".lower()) or {}
        sort_order = _walk_xmind_topic(root, [], raw_items, sort_order, is_root=True)
    normalized = normalize_test_points(raw_items)
    for idx, item in enumerate(normalized):
        item["sort_order"] = idx
    if not normalized:
        raise ValueError("XMind 中未识别到有效测试点（请确认叶子节点为测试点标题）")
    return normalized


def format_environment_hints_for_prompt(hints: Optional[dict]) -> str:
    if not hints or not isinstance(hints, dict):
        return "（未填写；请仅从需求正文摘录环境信息，勿编造）"
    labels = {
        "system_name": "被测系统",
        "test_env": "环境(SIT/UAT等)",
        "access_entry": "访问入口/URL",
        "deploy_version": "版本号",
        "client_requirements": "客户端/浏览器",
        "accounts_roles": "测试账号与角色",
        "test_data": "测试数据",
        "dependencies": "依赖服务",
        "tools": "工具",
    }
    lines = []
    for key, label in labels.items():
        val = hints.get(key)
        if not val:
            continue
        if isinstance(val, list):
            val = "；".join(str(v) for v in val if v)
        lines.append(f"- {label}：{val}")
    return "\n".join(lines) if lines else "（未填写；请仅从需求正文摘录）"


def normalize_scheme_content(content: dict, hints: Optional[dict] = None) -> dict:
    """归一化 LLM 方案 JSON，兼容旧版 environment 字段。"""
    if not isinstance(content, dict):
        return {}
    out = dict(content)
    env = out.get("environment")
    if not isinstance(env, dict):
        env = {}
    legacy_env = env.get("envs") or env.get("environment")
    legacy_acc = env.get("accounts")
    legacy_data = env.get("test_data")
    if legacy_env and not env.get("test_env"):
        env["test_env"] = str(legacy_env)
    if legacy_acc and not env.get("accounts_roles"):
        env["accounts_roles"] = legacy_acc if isinstance(legacy_acc, list) else [str(legacy_acc)]
    if legacy_data and not env.get("test_data"):
        env["test_data"] = str(legacy_data)
    for key in ("client_requirements", "accounts_roles", "dependencies", "tools"):
        if key in env and isinstance(env[key], str):
            env[key] = [env[key]] if env[key].strip() else []
    hints = hints or {}
    if isinstance(hints, dict):
        for key in (
            "system_name", "test_env", "access_entry", "deploy_version",
            "test_data", "client_requirements", "accounts_roles", "dependencies", "tools",
        ):
            hval = hints.get(key)
            if not hval:
                continue
            cur = env.get(key)
            empty = not cur or (isinstance(cur, list) and not cur) or cur == "待确认"
            if empty:
                env[key] = hval
    out["environment"] = env
    return out


def _render_env_section(env: dict) -> list[str]:
    lines = ["## 4. 测试环境与数据", ""]
    field_labels = [
        ("system_name", "被测系统"),
        ("test_env", "测试环境"),
        ("access_entry", "访问入口"),
        ("deploy_version", "版本/构建"),
        ("client_requirements", "客户端要求"),
        ("accounts_roles", "测试账号与角色"),
        ("test_data", "测试数据"),
        ("dependencies", "依赖与环境"),
        ("tools", "测试工具"),
    ]
    has_any = False
    for key, label in field_labels:
        val = env.get(key)
        if not val:
            continue
        has_any = True
        if isinstance(val, list):
            lines.append(f"### {label}")
            for item in val:
                lines.append(f"- {item}")
        else:
            lines.append(f"- **{label}**：{val}")
    if not has_any:
        lines.append("- （需求与测试点中未明确环境信息，执行前请与项目经理/运维确认）")
    lines.append("")
    return lines


def render_scheme_markdown(title: str, content: dict, test_points: list[dict]) -> str:
    """将结构化方案渲染为 Markdown。"""
    lines = [f"# {title}", ""]
    scope = content.get("scope") or {}
    if scope:
        lines.append("## 1. 测试范围")
        in_scope = scope.get("in_scope") or []
        out_scope = scope.get("out_of_scope") or []
        if in_scope:
            lines.append("### 范围内")
            for item in in_scope:
                lines.append(f"- {item}")
        if out_scope:
            lines.append("### 范围外")
            for item in out_scope:
                lines.append(f"- {item}")
        lines.append("")

    strategy = content.get("strategy") or {}
    if strategy:
        lines.append("## 2. 测试策略")
        for k, label in [
            ("types", "测试类型"),
            ("smoke_criteria", "冒烟准则"),
            ("regression_scope", "回归范围"),
            ("approach", "测试方法"),
        ]:
            val = strategy.get(k)
            if val:
                if isinstance(val, list):
                    lines.append(f"- **{label}**：{', '.join(str(v) for v in val)}")
                else:
                    lines.append(f"- **{label}**：{val}")
        lines.append("")

    lines.append("## 3. 测试点摘要")
    by_module: dict[str, list] = defaultdict(list)
    for pt in test_points:
        key = pt.get("module_path") or pt.get("main_module") or "未分类"
        by_module[key].append(pt)
    for mod, pts in sorted(by_module.items()):
        lines.append(f"### {mod}（{len(pts)} 项）")
        for pt in pts:
            pid = pt.get("id", "-")
            lines.append(
                f"- #{pid} [{pt.get('test_type', '')}] {pt.get('title', '')} "
                f"（{pt.get('priority', '')}）"
            )
        lines.append("")

    env = content.get("environment") or {}
    lines.extend(_render_env_section(env if isinstance(env, dict) else {}))

    exit_c = content.get("exit_criteria") or {}
    if exit_c:
        lines.append("## 5. 通过准则")
        for k, label in [("pass_rate", "通过率"), ("blockers", "阻塞条件"), ("deliverables", "交付物")]:
            val = exit_c.get(k)
            if val:
                if isinstance(val, list):
                    lines.append(f"- **{label}**：")
                    for item in val:
                        lines.append(f"  - {item}")
                else:
                    lines.append(f"- **{label}**：{val}")
        lines.append("")

    risks = content.get("risks") or []
    if risks:
        lines.append("## 6. 风险与假设")
        for r in risks:
            if isinstance(r, dict):
                lines.append(f"- **{r.get('item', '')}**：{r.get('mitigation', '')}")
            else:
                lines.append(f"- {r}")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def compute_overview(
    points: list[dict],
    schemes: list[dict],
    case_count: int = 0,
) -> dict:
    type_stats: dict[str, int] = defaultdict(int)
    module_stats: dict[str, int] = defaultdict(int)
    status_stats: dict[str, int] = defaultdict(int)
    for pt in points:
        type_stats[pt.get("test_type") or "其他"] += 1
        module_stats[pt.get("main_module") or "未分类"] += 1
        status_stats[pt.get("status") or "draft"] += 1
    latest_scheme = schemes[0] if schemes else None
    return {
        "test_point_total": len(points),
        "test_point_confirmed": status_stats.get("confirmed", 0),
        "test_point_draft": status_stats.get("draft", 0),
        "scheme_count": len(schemes),
        "latest_scheme": latest_scheme,
        "case_count": case_count,
        "type_distribution": dict(type_stats),
        "module_distribution": dict(module_stats),
    }


def format_test_points_for_prompt(points: list[dict], max_items: int = 80) -> str:
    if not points:
        return "（暂无测试点）"
    lines = []
    for i, pt in enumerate(points[:max_items], 1):
        lines.append(
            f"{i}. [#{pt.get('id', '-')}][{pt.get('test_type', '')}] "
            f"{pt.get('title', '')} | 模块:{pt.get('module_path', '')} | "
            f"优先级:{pt.get('priority', '')}\n"
            f"   说明:{(pt.get('description') or '')[:200]}"
        )
    if len(points) > max_items:
        lines.append(f"... 另有 {len(points) - max_items} 条未列出")
    return "\n".join(lines)


def format_test_points_for_case_generation(points: list[dict], max_items: int = 120) -> str:
    """供「测试点 → 功能用例」Prompt 使用的详细清单。"""
    if not points:
        return "（暂无测试点）"
    lines = []
    for pt in points[:max_items]:
        pid = pt.get("id", "-")
        lines.append(
            f"### 测试点 #{pid}\n"
            f"- 标题：{pt.get('title', '')}\n"
            f"- 类型：{pt.get('test_type', '')} | 优先级：{pt.get('priority', '')}\n"
            f"- 模块：{pt.get('module_path') or pt.get('main_module') or '未分类'}"
            f"{('/' + pt['sub_module']) if pt.get('sub_module') else ''}\n"
            f"- 说明：{(pt.get('description') or '（无）')[:400]}\n"
            f"- 验收参考：{(pt.get('acceptance_ref') or '（无）')[:200]}"
        )
    if len(points) > max_items:
        lines.append(f"\n... 另有 {len(points) - max_items} 条测试点未列出")
    return "\n\n".join(lines)


def test_points_cases_source_ref(batch_name: str = "") -> str:
    """功能用例 source_ref，与需求文档 batch: 区分。"""
    name = (batch_name or "").strip()
    return f"test_points:batch:{name}" if name else "test_points"


def _strip_md_bold(text: str) -> str:
    """标题等纯文本场景去掉 ** 标记。"""
    import re
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text or "")


def _add_rich_text(paragraph, text: str) -> None:
    """将含 **加粗** 的行写入 Word 段落（多 run）。"""
    if not text:
        return
    i = 0
    while i < len(text):
        start = text.find("**", i)
        if start == -1:
            paragraph.add_run(text[i:])
            break
        if start > i:
            paragraph.add_run(text[i:start])
        end = text.find("**", start + 2)
        if end == -1:
            paragraph.add_run(text[start:])
            break
        run = paragraph.add_run(text[start + 2 : end])
        run.bold = True
        i = end + 2


def _add_md_paragraph(doc, text: str, style: str | None = None):
    """新增段落并解析行内加粗。"""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    _add_rich_text(p, text)
    return p


def markdown_to_docx_bytes(content_md: str, title: str = "") -> bytes:
    """将 Markdown 方案正文转为 Word docx。"""
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(_strip_md_bold(title), 0)
    for line in (content_md or "").splitlines():
        text = line.rstrip()
        if not text:
            continue
        if text.startswith("# "):
            doc.add_heading(_strip_md_bold(text[2:].strip()), 1)
        elif text.startswith("## "):
            doc.add_heading(_strip_md_bold(text[3:].strip()), 2)
        elif text.startswith("### "):
            doc.add_heading(_strip_md_bold(text[4:].strip()), 3)
        elif text.startswith("  - "):
            _add_md_paragraph(doc, text[4:].strip(), style="List Bullet")
        elif text.startswith("- "):
            _add_md_paragraph(doc, text[2:].strip(), style="List Bullet")
        else:
            _add_md_paragraph(doc, text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
