"""测试文件在线预览（文本 / Office 转 HTML）"""
from __future__ import annotations

import html
from io import BytesIO
from typing import Optional


_PREVIEW_TEXT_EXT = {".txt", ".md", ".csv", ".json", ".pem", ".cer", ".crt", ".key"}
_PREVIEW_DOCX_EXT = {".docx", ".doc"}


def can_preview_filename(filename: str) -> bool:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return True
    if any(name.endswith(ext) for ext in _PREVIEW_TEXT_EXT):
        return True
    if any(name.endswith(ext) for ext in _PREVIEW_DOCX_EXT):
        return True
    return name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp"))


def preview_content_type(filename: str) -> Optional[str]:
    """返回 preview 响应类型：pdf | image | text | html | unsupported"""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
        return "image"
    if any(name.endswith(ext) for ext in _PREVIEW_TEXT_EXT):
        return "text"
    if any(name.endswith(ext) for ext in _PREVIEW_DOCX_EXT):
        return "html"
    return None


def _docx_to_html(content: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(f"<p>{html.escape(text)}</p>")
    for table in doc.tables:
        rows_html = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape((cell.text or '').strip())}</td>" for cell in row.cells)
            rows_html.append(f"<tr>{cells}</tr>")
        if rows_html:
            parts.append(f"<table>{''.join(rows_html)}</table>")
    return "".join(parts) if parts else "<p>（文档无文本内容）</p>"


def build_preview_payload(content: bytes, filename: str, mime_type: str = "") -> dict:
    kind = preview_content_type(filename)
    if kind == "text":
        text = content.decode("utf-8", errors="replace")
        return {"mode": "text", "content": text, "file_name": filename}
    if kind == "html":
        return {"mode": "html", "content": _docx_to_html(content), "file_name": filename}
    if kind in ("pdf", "image"):
        return {"mode": kind, "file_name": filename, "mime_type": mime_type}
    raise ValueError("该文件类型不支持在线预览")
